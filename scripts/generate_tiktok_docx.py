import docx
from docx.oxml.ns import qn
import re

template_path = r'd:\桌面\计算机\02\CV\vibe-cv-copilot\resources\排版范例.docx'
output_path = r'd:\桌面\计算机\02\CV\vibe-cv-copilot\output\20260517-TikTok_AI产品实习生-简历-沈妍-2028届-华东师范大学-18917534871.docx'

doc = docx.Document(template_path)

# 全局替换实习 Title
for para in doc.paragraphs:
    if "LLM算法实习生" in para.text:
        for run in para.runs:
            run.text = run.text.replace("LLM算法实习生", "AI策略产品实习生")

replacements = [
    # 实习1：枢宇科技
    ("数据工程 为探讨LLM量化决策的非理性行为，结合智能体工作流、联网搜索批量构建加密货币、商品期货、美股(SPY)跨市场情绪数据集，生成300个高保真场景；引入自动化数据质检门禁机制对抗流水线漏斗效应，拦截失衡数据；构建了包含多维特征的场景级宽表（Wide Table）作为底层数据基座", 
     "数据工程", "：利用 AutoGen 智能体工作流（Search Agent 检索新闻，Critic Agent 过滤噪声）自动化构建跨市场（Crypto/SPY/Futures）情绪数据集（300+ 高保真场景），并沉淀出包含多维特征的场景级 Wide Table；引入自动化数据质检门禁，成功对抗流水线漏斗效应，为后续算法评测提供高质量标准语料"),
    
    ("评测建模 引入行为金融学理论，构建 AI 量化行为偏差 Benchmark，部署并发引擎对26款 主流模型展开控制变量实验，得出共情能力更强的商业模型更易受噪音干扰而追涨杀跌；深度调研GitHub上20+款开源顶流金融Agent，验证了业界普遍采用的一刀切硬规则风控拦截在平淡场景下会造成极高误伤的痛点",
     "模型调优", "：为了解决 LLM 决策中的非理性痛点，协同算法团队构建 AI 量化行为偏差 Benchmark，对 30 款主流模型展开 Robustness Testing，精准量化了闭源商业模型的认知脆弱性（平均反转率高达 45.8%），并验证了代码特化模型在极端行情下的抗恐慌定力（反转率低至 34.9%），为风控策略优化指明方向"),
    
    ("架构设计 引入System2认知框架，利用宽表特征训练逻辑回归分类器作为第一层动态门控，仅在判定为高危场景时拉起多模型专家委员会进行多数票仲裁，在成功压低1.5%极端反转率的同时，实现正常场景 0 误伤",
     "策略规划", "：为了降低一刀切风控的高误伤率，主导设计 System 2 动态门控与多模型专家投票架构，在压低 1.5% 极端错误率的同时实现正常场景 0 误伤，完成从策略设计到产品落地的完整链路"),

    # 项目1：大创 (替换院企合作)
    ("基于多源数据融合的国产二次元IP热度评估指标构建研究（院企合作项目）",
     "多模态大模型幻觉测评系统（大创项目）", ""),
    
    ("2025.6-2025.12",
     "2024.4-至今", ""),
    
    ("项目概述 聚焦二游热度数据碎片化、动态变化等评估难点，设计热度指标体系，量化评估IP价值",
     "生态构建", "：深入拆解大模型在共现偏差与语言先验下的缺陷场景，引入主动陷阱设计的产品理念，运用图像编辑技术与 Prompt 工程，定向生成涵盖关系缺失等四大维度的 1600+ 诱导性对抗数据 (Golden Dataset)，探明大模型在极端内容场景下的能力边界"),
    
    ("数据爬取 用Python批量爬取脚本，获取143个国产二次元游戏IP在小红书、B站等多元异构 平台上的互动量数据",
     "链路设计", "：从 0 到 1 主导开发基于 Streamlit 的人工在环 (HITL) 数据清洗审核平台，完成从需求拆解、多模态非对称交互设计（无缝翻页、快捷键监听）到上线的全流程，实现海量图文数据与 VQA 问答对的极速即时纠错，大幅提升数据生产效率"),
    
    ("数据处理 针对多平台搜索算法差异导致的时间采样偏差，借助RFM模型构建时间衰减权重函数，最终在保留各平台数据特性的前提下实现热度指标的横向可比，为后续建模提供标准化数据基础",
     "算法协同", "：为深度剖析逻辑推演中的注意力偏移效应，创新提出 CoT 负优化率等动态评估指标，将黑盒的模型幻觉成因量化为可追踪的迭代数据，并规划基于此指标为后续对比解码、DPO 微调等算法定向调优提供数据牵引与方向"),

    # 项目2：数学建模
    ("项目概述 构建无人机投放烟幕弹的协同优化模型，设计多机多弹干扰策略以最大化来袭导 弹的累计遮蔽时长",
     "项目统筹", "：作为团队负责人统筹建模与开发进度，在 72 小时极限高压环境下完成业务问题拆解、技术路线选型与任务分发，确保项目高标准交付，最终斩获全国一等奖"),
    
    ("数学建模 建立多机多弹情形下遮蔽时长判定模型，将复杂三维运动中的轨迹预测误差控制 在1.5%以内",
     "策略规划", "：主导核心建模思路的讨论与决策，多维度对比多条技术路线的优劣并敲定最优解，有效协调算法手与编程手的跨通道合作，大幅降低沟通成本与返工率"),
    
    ("算法优化 主导求解路线比选，设计带有基因保留与早停机制的改进遗传算法。相较于传统 遗传算法，高维非凸问题的收敛速度提升约65%，单次求解耗时缩短至2.5秒内，且可行解搜索成功率达 98.5% 以上",
     "成果交付", "：用Latex将复杂的数据模型与代码逻辑转化为逻辑严密、结构清晰的学术论文，直击评委核心诉求，展现出极强的业务理解与跨专业书面表达能力"),

    # 项目3：自媒体博主
    ("账号定位 上海名校学习博主，专注本地K12教育及大学深造赛道，分享学习心得与技巧，提供个性化解决方案",
     "话题策划", "：在本地 K12 及深造赛道建立个人 IP，独立打通从选题策划到内容矩阵全链路，累计获赞与收藏 2w+，1w+ 浏览量爆款笔记 2 条，全网粉丝 3000+"),
    
    ("结果展示 累计获赞与收藏 2w+，1w+浏览量爆款笔记 2 条，全网粉丝3000+，公域转私域CVR达8%，搭建可落地知识付费商业模式，2 年内创收30W+，私域付费成交转化率40%",
     "用户增长", "：搭建可落地知识付费商业模式，公域转私域 CVR 达 8%，2 年内成交额 40W+，私域付费成交转化率 60%，实现流量到商业价值的闭环")
]

# Words that need to be bolded according to the visual priority rule (Numbers > English > Chinese)
# Note: In word documents, runs can have individual bolding.
words_to_bold = [
    "AutoGen", "Search Agent", "Critic Agent", "Crypto/SPY/Futures", "300+", "Wide Table",
    "LLM", "AI", "Benchmark", "30", "Robustness Testing", "45.8%", "34.9%",
    "System 2", "1.5%", "0",
    "Prompt", "1600+", "Golden Dataset", "0", "1", "Streamlit", "HITL", "VQA", "CoT", "DPO",
    "72", "Latex",
    "K12", "IP", "2w+", "1w+", "2", "3000+", "CVR", "8%", "40W+", "60%"
]

def clean_text(text):
    return re.sub(r'\s+', '', text)

# Map clean original text to (bold_header, content)
rep_map = {clean_text(orig): (header, content) for orig, header, content in replacements}

for para in doc.paragraphs:
    p_clean = clean_text(para.text)
    if not p_clean:
        continue

    for orig_clean, (header, content) in rep_map.items():
        if orig_clean in p_clean:
            if header == "" and content == "":
                # Change the whole paragraph to the new header
                para.clear()
                run = para.add_run(header)
                run.font.name = '楷体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')    
                run.font.size = docx.shared.Pt(10)
                break

            # Clear existing runs
            para.clear()

            # Add bold header if exists
            if header:
                run_bold = para.add_run(header)
                run_bold.bold = True
                run_bold.font.name = '楷体'
                run_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')    
                run_bold.font.size = docx.shared.Pt(10)

            # Add content (split by words_to_bold to apply bold formatting)
            # Create a regex pattern to match any of the words to bold
            # Sort by length descending to match longer strings first (e.g. 1w+ before 1)
            escaped_words = [re.escape(w) for w in sorted(words_to_bold, key=len, reverse=True)]
            pattern = re.compile(r'(' + '|'.join(escaped_words) + r')')
            
            parts = pattern.split(content)
            for part in parts:
                if not part:
                    continue
                run = para.add_run(part)
                run.font.name = '楷体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')     
                run.font.size = docx.shared.Pt(10)
                if part in words_to_bold:
                    run.bold = True
                    # Also set Times New Roman for english/numbers if needed, but keeping 楷体 is fine as per template
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')

            break

doc.save(output_path)
print(f"Successfully generated {output_path}")
